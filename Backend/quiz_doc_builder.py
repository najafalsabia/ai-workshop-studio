"""
Turns a generated quiz (from quiz_generator's generate_quiz) into a real
.docx file — questions and answers together, ready to open in Word.

No LLM calls here — pure post-processing, same role as notebook_builder.py
for labs (Step 4) and Step 8's python-pptx export for slides.

Uses python-docx (not the Node docx-js tool) since the rest of this
project is plain Python — no need for a second language/runtime just for
this one file.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

CORRECT_ANSWER_COLOR = RGBColor(0x1E, 0x7A, 0x34)  # a readable dark green


def save_quiz_docx(quiz_result: dict, output_path: str = "quiz.docx") -> str:
    """
    Takes generate_quiz's output ({"quiz": {"title": ..., "questions": [...]}})
    and writes ONE real .docx file with every question, its 4 options, and
    the correct answer marked inline — questions and answers together in
    the same file, per how this project's quizzes are meant to be
    reviewed (not two separate trainee/answer-key files, unlike the labs).

    Returns the file path written.
    """
    quiz = quiz_result["quiz"]
    doc = Document()

    title = doc.add_heading(quiz.get("title", "Quiz"), level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(f"{len(quiz.get('questions', []))} questions — answer key included")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True

    for i, q in enumerate(quiz.get("questions", []), start=1):
        doc.add_paragraph()  # spacing before each question

        question_para = doc.add_paragraph()
        number_run = question_para.add_run(f"{i}. ")
        number_run.bold = True
        text_run = question_para.add_run(q.get("question", ""))
        text_run.bold = True
        difficulty_run = question_para.add_run(f"   [{q.get('difficulty', '?')}]")
        difficulty_run.italic = True
        difficulty_run.font.size = Pt(9)

        correct = q.get("correct_answer")
        for letter, option in zip("ABCD", q.get("options", [])):
            option_para = doc.add_paragraph(style="List Bullet")
            option_run = option_para.add_run(f"{letter}. {option}")
            if option == correct:
                option_run.bold = True
                option_run.font.color.rgb = CORRECT_ANSWER_COLOR
                option_para.add_run("  ✓ Correct answer").italic = True

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    # Self-test — builds a real .docx from mock quiz data (no API keys, no
    # network needed) and verifies the file both exists AND is a genuinely
    # readable Word document (re-opened with python-docx, not just "the
    # code ran without crashing").

    mock_quiz_result = {
        "quiz": {
            "title": "Beyond Syntax — Final Quiz",
            "questions": [
                {
                    "question": "What is the primary risk of ambiguous prompts?",
                    "options": ["Slower response time", "Unreliable output", "Higher cost", "Shorter answers"],
                    "correct_answer": "Unreliable output",
                    "difficulty": "easy",
                },
                {
                    "question": "Which step comes first in the Trust-But-Verify pipeline?",
                    "options": ["Manual audit", "Syntax linting", "Deployment", "Documentation"],
                    "correct_answer": "Syntax linting",
                    "difficulty": "medium",
                },
            ],
        }
    }

    print("=== Test 1: save_quiz_docx writes a real file ===")
    output_path = "/tmp/test_quiz_output.docx"
    result_path = save_quiz_docx(mock_quiz_result, output_path)
    print(f"saved to: {result_path}")

    import os
    assert os.path.isfile(result_path)
    print("file exists on disk: True")

    print("\n=== Test 2: the file is a genuinely readable .docx (re-opened with python-docx) ===")
    reopened = Document(result_path)
    paragraph_texts = [p.text for p in reopened.paragraphs if p.text.strip()]
    for line in paragraph_texts:
        print(" ", line)

    assert any("Beyond Syntax" in t for t in paragraph_texts)
    assert any("Unreliable output" in t for t in paragraph_texts)
    assert any("Correct answer" in t for t in paragraph_texts)
    print("\nAll self-tests passed.")
